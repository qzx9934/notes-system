# -*- coding: utf-8 -*-
"""生成「五小创新」成果报告 Word 文档：
基于 AI 大模型的工作笔记系统在电厂的应用。"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BODY_FONT = '宋体'
HEAD_FONT = '黑体'
TITLE_FONT = '黑体'
PRIMARY = RGBColor(0x1F, 0x4E, 0x79)


def set_cn_font(run, font=BODY_FONT, size=12, bold=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color
    r = run._element.rPr.rFonts
    r.set(qn('w:eastAsia'), font)
    r.set(qn('w:ascii'), font)
    r.set(qn('w:hAnsi'), font)


def para(doc, text='', size=12, font=BODY_FONT, bold=False, align=None,
         color=None, first_indent=True, space_after=6, line=22):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.line_spacing = Pt(line)
    if first_indent and align not in (WD_ALIGN_PARAGRAPH.CENTER,):
        pf.first_line_indent = Pt(size * 2)
    if text:
        run = p.add_run(text)
        set_cn_font(run, font=font, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level=1):
    sizes = {1: 15, 2: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run(text)
    set_cn_font(run, font=HEAD_FONT, size=sizes[level], bold=True,
                color=PRIMARY if level == 1 else None)
    return p


def bullet(doc, text, size=12):
    p = doc.add_paragraph(style=None)
    p.paragraph_format.left_indent = Pt(24)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(22)
    run = p.add_run('● ' + text)
    set_cn_font(run, size=size)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def make_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = c.paragraphs[0].add_run(h)
        set_cn_font(run, font=HEAD_FONT, size=11, bold=True)
        shade_cell(c, 'D9E2F3')
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            run = p.add_run(val)
            set_cn_font(run, size=10.5)
            p.paragraph_format.line_spacing = Pt(18)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def build():
    doc = Document()
    # 默认样式
    normal = doc.styles['Normal']
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn('w:eastAsia'), BODY_FONT)
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(2.8); s.right_margin = Cm(2.8)

    # ---------- 封面 ----------
    for _ in range(3):
        doc.add_paragraph()
    para(doc, '“五小”创新成果报告', size=16, font=HEAD_FONT, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_after=4)
    para(doc, '（小革新·小发明类）', size=12, align=WD_ALIGN_PARAGRAPH.CENTER,
         first_indent=False, space_after=30)
    para(doc, '基于 AI 大模型的工作笔记系统', size=22, font=TITLE_FONT, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY, first_indent=False, space_after=4)
    para(doc, '在电厂的应用', size=22, font=TITLE_FONT, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=PRIMARY, first_indent=False, space_after=40)
    for _ in range(4):
        doc.add_paragraph()
    info = [('成果名称', '基于 AI 大模型的工作笔记系统在电厂的应用'),
            ('成果类别', '“五小”创新 —— 小革新 / 小发明'),
            ('专业领域', '集控运行 / 信息化'),
            ('完 成 人', '＿＿＿＿＿＿＿＿＿＿＿'),
            ('完成单位', '＿＿＿＿＿＿＿＿＿＿＿'),
            ('完成日期', '二〇二六年五月')]
    t = doc.add_table(rows=0, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in info:
        cells = t.add_row().cells
        cells[0].width = Cm(3.2); cells[1].width = Cm(9)
        rk = cells[0].paragraphs[0].add_run(k)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cn_font(rk, font=HEAD_FONT, size=12, bold=True)
        rv = cells[1].paragraphs[0].add_run(v)
        set_cn_font(rv, size=12)
    doc.add_page_break()

    # ---------- 摘要 ----------
    heading(doc, '摘  要', 1)
    para(doc, '针对电厂集控运行专业知识分散、经验难以沉淀、新员工培养周期长等长期痛点，'
              '本成果自主设计并开发了一套“基于 AI 大模型的工作笔记系统”。系统以 5 大领域、'
              '26 个子类的标准化知识编码体系为骨架，融合全文检索、图片、Markdown、收藏、'
              '卡片漫游等实用功能。系统创新性地接入 DeepSeek 大模型，在网页端实现笔记的 AI 总结、'
              'AI 填充、AI 整理与一键批量处理；更进一步，配套整理提示词与 AI Agent Skill，'
              '可让智能体把整篇 Word/PDF/PPTX 资料自动整理成规范笔记并经 API 批量入库，'
              '实现“资料→知识库”的全自动流水线。同时设计了管理员 / 共建者 / 查看者三级权限与“提交—审批”'
              '协作机制，做到人人可补充、质量可管控。系统已在浏览器与手机端（PWA）投入使用，'
              '显著提升了运行知识的沉淀效率与共享水平，对保障机组安全运行、加速人才培养具有'
              '良好的应用与推广价值。')
    para(doc, '关键词：人工智能；大语言模型；知识管理；集控运行；安全生产；数字化', bold=True)

    # ---------- 一、背景 ----------
    heading(doc, '一、立项背景与现状问题', 1)
    para(doc, '电厂集控运行是典型的知识密集型、高风险岗位，运行人员日常需要掌握大量规程要点、'
              '操作经验、保护定值、事故预案与经验反馈。长期以来，这些宝贵的知识与经验主要存在'
              '以下问题：')
    bullet(doc, '一是“散”。资料散落在纸质笔记、各类文档、聊天记录与个人脑海中，缺乏统一归集，查找困难。')
    bullet(doc, '二是“失”。骨干、老师傅的经验高度依赖个人，人员流动或退休即造成知识流失，难以传承。')
    bullet(doc, '三是“慢”。新员工面对海量规程无从下手，缺乏体系化、可检索的学习载体，培养周期长。')
    bullet(doc, '四是“乱”。即便有人整理，格式不统一、口径不一致，且夹杂大量排班通知等事务性信息，质量参差。')
    para(doc, '在班组日常工作与人才培养的双重需求下，迫切需要一个低成本、易上手、可持续积累，'
              '并能借助人工智能减轻整理负担的知识管理工具。这正是本成果的立项初衷。')

    # ---------- 二、成果简介 ----------
    heading(doc, '二、成果简介与创新目标', 1)
    para(doc, '本成果是一套面向电厂运行人员的 Web 工作笔记管理系统，采用轻量化的 Flask + SQLite '
              '技术栈，浏览器即可使用，并支持手机端安装为 App（PWA）。其核心目标可概括为'
              '“一个体系、两套智能、两级协作”：')
    bullet(doc, '一个体系：建立 5 大领域、26 子类的电厂运行知识标准编码，笔记自动编号、可分级分类。')
    bullet(doc, 'AI Agent 全自动整理上传：配套提示词与 Skill，让 AI 智能体把整篇资料（Word/PDF/PPTX 等）'
                '自动抽取、整理成规范笔记并经 API 批量入库，实现“资料→知识库”的流水线作业。')
    bullet(doc, '系统内置在线智能：网页端提供 AI 总结、AI 填充、AI 整理，把单条笔记的整理交给 AI。')
    bullet(doc, '两级协作：管理员与共建者分级协同，写操作经审批入库，兼顾全员参与与内容质量。')

    # ---------- 三、技术方案 ----------
    heading(doc, '三、技术方案与系统架构', 1)
    para(doc, '系统采用前后端一体的轻量化架构，便于在电厂内网或云服务器低成本部署、稳定运行：')
    make_table(doc,
               ['层次', '技术选型与说明'],
               [['前端', '单页应用（HTML/CSS/JS），本地内置 marked / DOMPurify，无外部 CDN 依赖；'
                         '支持 PWA 离线外壳与“添加到主屏幕”。'],
                ['后端', 'Python Flask 提供 RESTful API；按角色鉴权，输入严格校验、内容渲染消毒。'],
                ['数据', 'SQLite 单文件数据库，启用 FTS5 全文索引（中文子串检索，自动回退 LIKE）；'
                         '图片按内容哈希去重存储。'],
                ['AI', '通过环境变量接入 DeepSeek 大模型，密钥不落代码、调用具备限流退避重试；'
                       '并配套提示词与 organize-notes Skill，支持 AI Agent 自动整理资料、经 ingest 接口批量上传。'],
                ['安全', '会话 Cookie HttpOnly、登录失败限流、安全响应头、HTTPS/反向代理加固、API 令牌。']],
               widths=[2.6, 10.2])
    para(doc, '后端共配套 90 余项自动化测试用例，覆盖鉴权、字段校验、编号生成、审批流、AI 接口等关键路径，'
              '保障迭代质量与运行可靠性。', space_after=6)

    # ---------- 四、主要功能与创新点 ----------
    heading(doc, '四、主要功能与创新点', 1)
    heading(doc, '（一）标准化知识体系，让经验“有处可归”', 2)
    para(doc, '将电厂运行知识划分为系统设备、运行操作、安全管理、技术标准、综合管理 5 大领域共 26 个子类，'
              '每条笔记归属唯一章节并自动生成形如 A01-001 的编号，配合重要等级（★/★★/★★★）与来源标注，'
              '形成结构清晰、可检索、可分级的知识库。')
    heading(doc, '（二）AI Agent + Skill + API，让资料“一句话”自动整理入库（核心创新）', 2)
    para(doc, '这是本成果最具特色的创新点：不再逐条手敲笔记，而是把一整套“整理—归类—上传”工作交给 '
              'AI 智能体（Agent）自动完成。为此设计了三件配套：标准化的整理提示词、可在 Claude Code 等'
              '工具中一键调用的 Skill（organize-notes，内含文档抽取与上传脚本），以及支持批量去重入库的 '
              'API 接口。整个流程形成“四步流水线”：')
    make_table(doc,
               ['步骤', '内容'],
               [['① 抽取文本', '对 Word / PDF / PPTX / Markdown 等各种格式资料，自动抽取出纯文本。'],
                ['② AI 整理', '智能体依据提示词把资料按 26 个章节归类，整理成符合字段规范（标题/标签/'
                              '等级/来源/正文）的结构化 JSON，原样保留定值、动作条件等关键数据。'],
                ['③ API 上传', '经 /api/notes/ingest 批量入库，按“章节+标题”自动去重：已存在则更新、'
                               '新内容则新增，可反复增量上传。'],
                ['④ 验证回填', '返回新增/合并/跳过明细，便于核对；手动插入的图片在更新时自动保留不丢失。']],
               widths=[2.2, 10.6])
    para(doc, '使用者只需在 AI 工具中说一句“把这些资料整理成笔记上传”，即可完成从原始文档到规范知识库的'
              '全过程，整理效率较人工录入实现数量级提升。整理用 API 令牌经环境变量传递、'
              '权限可控、可随时吊销，兼顾自动化与安全。')
    heading(doc, '（三）系统内置在线 AI，让单条整理“省时省力”', 2)
    para(doc, '在网页编辑场景中，系统把大模型能力直接嵌入操作界面：')
    bullet(doc, 'AI 总结：自动把一条笔记提炼为要点式 Markdown 摘要，便于快速回顾与考前复习。')
    bullet(doc, 'AI 填充：根据正文自动推断标题、标签、所属章节、重要等级与来源，录入“一键成型”。')
    bullet(doc, 'AI 整理：一次调用即可清理排班/通知等非知识性内容、规范乱序的小序号，并同步完成字段填充；'
                '改动克制、不篡改技术数值，且提供“撤回整理”保障安全。')
    bullet(doc, '一键批量总结：可对勾选的多条笔记批量生成摘要，遇限流自动退避重试，并反馈成功/跳过/失败明细。')
    bullet(doc, '提示词可配置：管理员可随时调整 AI 总结与 AI 整理的提示词，使输出贴合本厂术语与口径。')
    heading(doc, '（四）三级权限 + 审批协作，让共建“质量可控”（机制创新）', 2)
    para(doc, '设管理员、共建者、查看者三级角色。共建者的新增、编辑、删除不直接生效，而是生成“变更申请”，'
              '管理员在收件箱可清晰对比改动明细（原值→新值）后批准或驳回。既调动了全员积累知识的积极性，'
              '又通过审批关口守住了知识库的准确性与权威性。')
    heading(doc, '（五）实用体验细节，让使用“顺手好用”', 2)
    bullet(doc, '全文检索 + 多维筛选（领域/章节/等级/来源）+ 多种排序，秒级定位所需知识。')
    bullet(doc, '支持粘贴截图、拖拽上传图片与 Markdown 排版；详情页可就地快速改字段。')
    bullet(doc, '“笔记漫游”全屏卡片随机复习、翻面看 AI 摘要；每位用户拥有独立收藏夹。')
    bullet(doc, '登录日志与最后上线时间便于管理；命令行/大模型可经 API 令牌批量录入，无需打开浏览器。')

    # ---------- 五、实施应用 ----------
    heading(doc, '五、实施应用情况', 1)
    para(doc, '系统已完成开发、测试并投入试用，部署方式灵活：可运行于电厂内网服务器，亦可部署在云服务器'
              '经 HTTPS 对外提供，运行人员通过电脑浏览器或手机即可访问。班组可由管理员维护知识体系框架，'
              '运行人员以共建者身份在交接班、培训、事故分析后随手补充笔记，借助 AI 整理快速成稿、提交审批入库，'
              '逐步形成本班组、本专业的“活”知识库。')

    # ---------- 六、效益分析 ----------
    heading(doc, '六、应用效果与效益分析', 1)
    make_table(doc,
               ['效益类型', '具体体现'],
               [['安全效益', '规程要点、保护定值、事故预案随查随得，异常工况下辅助快速决策；'
                             '经验反馈系统化沉淀，减少同类问题重复发生，助力机组安全稳定运行。'],
                ['管理效益', '知识从“个人私藏”转为“班组共享”，骨干经验得以留存传承；'
                             '审批机制保证口径统一、内容权威，提升标准化管理水平。'],
                ['培训效益', '新员工有了体系化、可检索、带 AI 摘要的学习载体，'
                             '上手更快，缩短培养周期，降低带教成本。'],
                ['经济效益', '基于开源技术栈自主开发，部署运维成本低；'
                             'AI 自动整理大幅降低人工录入与编辑工时，把人力投向更有价值的分析工作。']],
               widths=[2.6, 10.2])
    para(doc, '综合来看，本成果以较小的投入，解决了运行知识“散、失、慢、乱”的痛点，'
              '在安全、管理、培训、经济等方面均产生了实实在在的效益。')

    # ---------- 七、推广价值 ----------
    heading(doc, '七、推广应用价值', 1)
    para(doc, '本成果具有良好的通用性与可复制性：知识编码体系可按不同专业（如检修、化学、燃料）灵活调整；'
              '技术栈轻量、部署简单、无外部依赖，便于在各班组、各车间乃至兄弟电厂推广；'
              'AI 提示词可配置，能快速适配不同单位的术语与管理口径。随着大模型能力的持续进步，'
              '系统在智能问答、知识图谱、隐患辨识等方向具备进一步拓展空间。')

    # ---------- 八、结论 ----------
    heading(doc, '八、结论与展望', 1)
    para(doc, '“基于 AI 大模型的工作笔记系统”立足班组一线实际需求，用信息化与人工智能手段，'
              '把分散易失的运行经验转化为结构化、可检索、可传承的数字资产，是“小革新”解决“大问题”的'
              '典型实践。下一步将结合使用反馈持续优化，探索 AI 智能问答、与生产实时数据联动等功能，'
              '让知识更好地服务于电厂安全生产与人才成长。')

    out = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
                       '基于AI大模型的工作笔记系统在电厂的应用-五小创新报告.docx')
    doc.save(out)
    print(out)


if __name__ == '__main__':
    build()
